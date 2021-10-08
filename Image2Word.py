import tkinter as tk                # GUI
import tkinter.filedialog           # Loading file/ image from file system

import docx                         # Handling Word documents
import cv2 as cv
import docx2txt                    # Image processing in particular table remove
import numpy as np                  
import pytesseract                  # Image to string conversion
import re                           # Searching for expresion in document

class Image2Word():

    def __init__(self, master):
        
        self.master = master
        master.title("Image2Word")

        # file path 
        self.open_file_name = tk.Label(master, text= "Temp")
        self.open_file_name.grid(row= 0 ,column= 0)

        # open button for word documents
        self.open_doc_button = tk.Button(master, text= "Open Document", command= self.open_document)
        self.open_doc_button.grid(row=2, column= 0)
        
        # open button for image
        self.open_img_button = tk.Button(master, text= "Open Image", command= self.open_image)
        self.open_img_button.grid(row= 4, column= 0)

        # convert button
        self.convert_button = tk.Button(master, text="\N{RIGHTWARDS BLACK ARROW}", command= self.image_to_text)
        self.convert_button.grid(row= 2, column= 2)

        # save button
        self.save_button = tk.Button(master, text= "Save to Docx", command= self.save_2_docs)
        self.save_button.grid(row= 2, column= 4)

        # close button
        self.close_button = tk.Button(master, text= "Close", command= master.quit)
        self.close_button.grid(row = 8, column= 2)

    def open_document(self):
        file_path = tkinter.filedialog.askopenfilename(
            filetypes = [("Word Document", "*.docx")]
        )
        if not file_path:
            return
        self.open_file_name['text'] = file_path
        self.temp_txt = docx2txt.process(file_path)
        self.document = docx.Document()
        

    def open_image(self):
        image_path = tkinter.filedialog.askopenfilename(
            filetypes=[("PNG", "*.png"), ("JPEG", "*.jpg")]
        )
        if not image_path:
            return
        self.img = cv.imread(image_path)

    def image_to_text(self):
        # tesseract command
        pytesseract.pytesseract.tesseract_cmd = r"D:\Tesseract-OCR\tesseract.exe"
        
        '''
            Conversion to gray scale, bitwise not operation on converted image 
            and applying of adaptive threshold to that image 
        '''
        gray = cv.cvtColor(self.img, cv.COLOR_BGR2GRAY)
        gray = cv.bitwise_not(gray)
        bw = cv.adaptiveThreshold(gray, 255, cv.ADAPTIVE_THRESH_MEAN_C, cv.THRESH_BINARY, 5, -3)

        # Copy of processed image to extract of horizontal and vertical lines
        horizontal = np.copy(bw)
        vertical = np.copy(bw)

        # [horizontal]
        # size on horizontal plane
        cols = horizontal.shape[1]
        horizontal_size = cols // 30

        # structure that keeps extracted horizontal lines thanks to morphologic operations
        horizontal_struct = cv.getStructuringElement(cv.MORPH_RECT, (horizontal_size, 1))

        # use of morphologic operations on copy of src image
        horizontal = cv.erode(horizontal, horizontal_struct, iterations= 3)
        horizontal = cv.dilate(horizontal, horizontal_struct, iterations= 3)
        # [/horizontal]

        # [vertical]
        # size on vertical plane
        rows = vertical.shape[0]
        vertical_size = rows // 20

        # structure that keeps extracted vertical lines thanks to morphologic operations
        vertical_struct = cv.getStructuringElement(cv.MORPH_RECT, (1, vertical_size))

        # use of morphologic operations on copy of src image
        vertical = cv.erode(vertical, vertical_struct)
        vertical = cv.dilate(vertical, vertical_struct)
        # [/vertical]

        # removing of all boxes 
        result = cv.subtract(bw, vertical + horizontal)
        result = cv.bitwise_not(result)

        '''
            Extract edges and smooth image according to the logic
            1. extract edges
            2. dilate(edges)
            3. src.copyTo(smooth)
            4. blur smooth img
            5. smooth.copyTo(src, edges)
        '''

        # Step 1
        edges = cv.threshold(result, 0, 255, cv.THRESH_BINARY | cv.THRESH_OTSU)[1]

        # Step 2
        kernel = np.ones((3,3), np.uint8)
        edges = cv.dilate(edges, kernel, iterations= 3)

        # Step 3
        smooth = np.copy(result)

        # Step 4
        smooth = cv.blur(smooth, (2,2))

        # Step 5
        (rows, cols) = np.where(edges != 0)
        result[rows, cols] = smooth[rows, cols]

        # converts processed image to string --- in future add more languages for now pol is ok
        custom = r'-l jpn+pol --psm 6'
        self.data = pytesseract.image_to_string(result,config = custom)
        self.data = re.sub("  +", " - ", self.data)

    def save_2_docs(self):
        splited = self.data.splitlines()

        to_add = []
        for line in splited:
            print(line)
            if re.search(line, self.temp_txt):
                continue
            else:
                to_add.append(line)
                self.temp_txt += line + "\n"

        
        para = self.document.add_paragraph()
        for line in to_add:
            para.add_run(line + "\n")
        
        to_add.clear()
        self.document.save(self.open_file_name['text'])
        #self.temp_txt = docx2txt.process(self.open_file_name['text'])
                       

def main():
    root = tk.Tk()
    my_gui = Image2Word(root)
    root.mainloop()


if __name__ == "__main__":
    main()